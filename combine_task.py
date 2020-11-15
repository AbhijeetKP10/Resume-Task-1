import docx
import os
import spacy
nlp=spacy.load('en_core_web_sm')
os.chdir(r'F:\BACKUP FOLDER\Internship')
doc=docx.Document('abhi.docx')
table = doc.tables[0]
z=doc.sections[0]
header=z.header
p=len(doc.paragraphs)
r=len(header.paragraphs)
data=[]
fg=0
rep=[]
d=os.stat("doc_to_txt.txt").st_size == 0
if(d==False):
    print("***********\nText file which contain the data is not empty.The file should not contain data. Please delete the data and then try again...\n***********")
else:
    print("***********\nThe selected file is empty. Now showing the result.\n***********")
#header
    for x in range (r):
        a=header.paragraphs[x].text
        if a is not '':
            with open('doc_to_txt.txt','a') as file:
                file.write(a+'\n')
            file.close()
#paragra
    for x in range (p):
        a=doc.paragraphs[x].text
        if a is not '':
            with open('doc_to_txt.txt','a') as file:
                file.write(a+'\n')
            file.close()
    with open('doc_to_txt.txt','a') as file:
        file.write('Table content:\n')
    file.close()
#table
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        content = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(content)
            continue
        row= dict(zip(keys, content))
        row_str=str(row)
        with open('doc_to_txt.txt','a') as file:
            file.write(row_str+'\n')
        file.close()
        data.append(row)
#finding the repeated words
    f=open('doc_to_txt.txt','r')
    rep=[]
    w=[]
    check=[]
    check.clear()
    data.clear()
    rep.clear()
    w.clear()
    ct=0
    flag=0
    for a in f.readlines():
        data.append(a.replace('\n','').replace("'","").replace(',','').replace(':','').replace(';','').replace('\t','').split(' '))
    f.close()
    for i in range(0,len(data)):
        for j in range(0,len(data[i])):
            check.append(data[i][j])
    for x in range(0,len(check)):
        z=check[x]
        if(x==0):
            rep.append(check[x])
        for y in range(0, len(rep)):
            if(rep[y]==check[x]):
                flag=1
                break
        if(flag==0):
            rep.append(z)
        flag=0
    for i in range(0,len(rep)):
        for j in range(0,len(check)):
            if(rep[i]==check[j]):
                ct=ct+1
        w.append(ct)
        ct=0
    
    for p in range (0,len(rep)):
        if rep[p]!='':
            with open('rep_count.txt','a') as file:
                file.write(rep[p]+'\t'+str(w[p])+'\n')
            file.close()
            
print(rep)
#####################################################
f=open("doc_to_txt.txt","r")
f1=f.readlines()
a=''
for x in f1:
    a=a+x
print(a)
nlp_a=nlp(a)
num=0
name=[]
email=''
num=''
noun=[]
flag=0
b=list(nlp_a.sents)
##############
#print(len(b))
count=0
a=[]
for ent in nlp_a.ents:
    count=+1
    if(ent.label_=='PERSON'):
        a=str(ent.text)
        name.append(a)
        fg=1
    if(count==1):
        break
    if(fg==1):
        break
for i in range(0,6):
    for word in list(nlp_a.sents)[i]:
        if(word.pos_=='NUM'):
            #print(word)
            if(len(word)==10):
                    try:
                        num=int(str(word))
                        flag=flag+1
                        break
                    except ValueError:
                        print('')
        elif(word.pos_=='PROPN' or word.pos_=='NOUN'):
            if(fg==1):
                break
            else:
                name.append(str(word))
        elif(word.pos_=='X'):
            email=word
            flag=flag+1
            
    if(flag==2):
        break
a=len(email)    
print(a)
if(a==0):
    for i in rep:
        for j in i:
            if(j=='@'):
                print("1")
                email=i
                break
q=['RESUME','Resume','of','Personal','Information','Name','Date','Miss.','No','Details','2nd','May','Mobile','Number','Birth','emailid','EMAILID','Emailid','Email','id','EMAIL','ID','Contact','CONTACT','contact','NO','no','November','PERSONAL','','INFORMATION']
for i in q:
    try:
        name.remove(i)
    except ValueError:
        continue
name = list(dict.fromkeys(name))
print("Name of the Candidate:",end='')
for i in range (len (name)):
    print("{} ".format(name[i]),end='')
print("\nMobile Number:",num)
print("Email id: ",email)       
##################
#r=len(list(nlp_dt.sents))